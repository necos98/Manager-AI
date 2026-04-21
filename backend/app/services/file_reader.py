from __future__ import annotations

import logging
import os

logger = logging.getLogger(__name__)

MAX_CHARS = 500_000
MAX_ROWS_PER_SHEET = 1000
LEGACY_FORMATS = {"doc", "xls"}


class ExtractionResult:
    __slots__ = ("text", "status", "error")

    def __init__(self, text: str, status: str, error: str | None = None):
        self.text = text
        self.status = status
        self.error = error


def extract(path: str, ext: str) -> ExtractionResult:
    ext = ext.lower().lstrip(".")
    if ext in LEGACY_FORMATS:
        return ExtractionResult("", "unsupported", f"Legacy format '.{ext}' not supported. Convert to .{ext}x.")
    try:
        if ext in ("txt", "md"):
            text = _extract_text(path)
        elif ext == "pdf":
            text = _extract_pdf(path)
        elif ext == "docx":
            text = _extract_docx(path)
        elif ext == "xlsx":
            text = _extract_xlsx(path)
        else:
            return ExtractionResult("", "unsupported", f"Extension '.{ext}' has no parser")
    except ImportError as e:
        logger.warning("Parser missing for %s: %s", ext, e)
        return ExtractionResult("", "failed", f"Parser library missing: {e}")
    except Exception as e:
        logger.exception("Extraction failed for %s", path)
        return ExtractionResult("", "failed", str(e))

    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]
    return ExtractionResult(text, "ok", None)


def _extract_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    parts: list[str] = []
    total = 0
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception as e:
            logger.debug("PDF page extract failed: %s", e)
            t = ""
        parts.append(t)
        total += len(t)
        if total > MAX_CHARS:
            break
    return "\n\n".join(parts)


def _extract_docx(path: str) -> str:
    import docx

    doc = docx.Document(path)
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(path: str) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    total = 0
    for sheet in wb.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i >= MAX_ROWS_PER_SHEET:
                parts.append(f"... (truncated at {MAX_ROWS_PER_SHEET} rows)")
                break
            cells = ["" if v is None else str(v) for v in row]
            if any(cells):
                line = " | ".join(cells)
                parts.append(line)
                total += len(line)
                if total > MAX_CHARS:
                    break
        if total > MAX_CHARS:
            break
    wb.close()
    return "\n".join(parts)


def file_is_low_text(text: str, file_size_bytes: int, threshold_chars: int = 50, threshold_bytes: int = 100_000) -> bool:
    return len(text) < threshold_chars and file_size_bytes > threshold_bytes
