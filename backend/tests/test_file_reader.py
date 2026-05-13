import os
import tempfile

import pytest

from app.services import file_reader


def test_extract_txt(tmp_path):
    p = tmp_path / "note.txt"
    p.write_text("hello world\nsecond line", encoding="utf-8")
    r = file_reader.extract(str(p), "txt")
    assert r.status == "ok"
    assert "hello world" in r.text
    assert "second line" in r.text


def test_extract_md(tmp_path):
    p = tmp_path / "doc.md"
    p.write_text("# Title\n\nparagraph", encoding="utf-8")
    r = file_reader.extract(str(p), "md")
    assert r.status == "ok"
    assert "Title" in r.text


def test_legacy_unsupported(tmp_path):
    p = tmp_path / "old.doc"
    p.write_bytes(b"fake")
    r = file_reader.extract(str(p), "doc")
    assert r.status == "unsupported"
    assert "Legacy" in r.error


def test_extract_docx(tmp_path):
    import docx
    doc = docx.Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph with keyword QUERYABLE")
    p = tmp_path / "d.docx"
    doc.save(str(p))
    r = file_reader.extract(str(p), "docx")
    assert r.status == "ok"
    assert "QUERYABLE" in r.text


def test_extract_xlsx(tmp_path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["col_a", "col_b"])
    ws.append([1, "hello"])
    p = tmp_path / "s.xlsx"
    wb.save(str(p))
    r = file_reader.extract(str(p), "xlsx")
    assert r.status == "ok"
    assert "Sheet1" in r.text
    assert "hello" in r.text


def test_extract_pdf(tmp_path):
    try:
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf not installed")
    try:
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not installed for test pdf generation")
    p = tmp_path / "t.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(100, 750, "PDF TEST CONTENT")
    c.save()
    r = file_reader.extract(str(p), "pdf")
    assert r.status == "ok"
    assert "PDF TEST CONTENT" in r.text


def test_truncation(tmp_path, monkeypatch):
    monkeypatch.setattr(file_reader, "MAX_CHARS", 50)
    p = tmp_path / "big.txt"
    p.write_text("x" * 200, encoding="utf-8")
    r = file_reader.extract(str(p), "txt")
    assert r.status == "ok"
    assert len(r.text) == 50


def test_audio_extensions_defined():
    assert "ogg" in file_reader.AUDIO_EXTENSIONS
    assert "mp3" in file_reader.AUDIO_EXTENSIONS
    assert "wav" in file_reader.AUDIO_EXTENSIONS


def test_extract_audio_dispatcher_routes(tmp_path, monkeypatch):
    def fake_extract_audio(path):
        return "transcribed text"

    monkeypatch.setattr(file_reader, "_extract_audio", fake_extract_audio)
    p = tmp_path / "audio.mp3"
    p.write_bytes(b"fake audio data")
    r = file_reader.extract(str(p), "mp3")
    assert r.status == "ok"
    assert r.text == "transcribed text"


def test_extract_audio_missing_whisper(tmp_path, monkeypatch):
    def raise_import_error():
        raise ImportError("No module named 'whisper'")

    monkeypatch.setattr(file_reader, "_get_whisper_model", raise_import_error)
    p = tmp_path / "audio.ogg"
    p.write_bytes(b"fake ogg data")
    r = file_reader.extract(str(p), "ogg")
    assert r.status == "failed"
    assert "whisper" in r.error.lower()


def test_extract_audio_unsupported_format(tmp_path):
    p = tmp_path / "unknown.xyz"
    p.write_bytes(b"unknown")
    r = file_reader.extract(str(p), "xyz")
    assert r.status == "unsupported"
